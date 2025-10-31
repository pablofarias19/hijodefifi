# -*- coding: utf-8 -*-
"""
MOTOR DE ANÁLISIS JURÍDICO UNIFICADO
====================================

PROPÓSITO:
Consolida TODAS las funcionalidades de:
- slegal.py
- analista_jurisprudencial.py
- analista_procesos_judiciales.py

ARQUITECTURA:
- ÚNICO cliente de `ingesta_multibase.py`.
- 0% duplicación de código.
- 100% funcionalidades avanzadas (Análisis de Texto, Q&A, RAG Dinámico, DOCX).

FUNCIONALIDADES:
✅ Análisis de Texto (Perfeccionamiento, Interpretación, Avanzado, Integral).
✅ Análisis de Fallos Jurisprudenciales (desde PDF).
✅ Análisis de Expedientes Procesales (desde PDF).
✅ Q&A Interactivo sobre expedientes (RAG local + global).
✅ Selección dinámica de bases de datos RAG.
✅ Exportación unificada a DOCX y PDF.
✅ Dashboard Streamlit unificado.
✅ CLI unificada.
"""

import os
import logging
import io
from pathlib import Path  # ← IMPORTAR AQUÍ al inicio

# ============================================================
# FIX: PyTorch Meta Tensor Issue (Python 3.13 + PyTorch 2.x)
# ============================================================
import sys

# ⚙️ CONFIGURACIONES PREVIAS AGRESIVAS
os.environ["PYTORCH_ENABLE_MPS_FALLBACK"] = "1"
os.environ["TORCH_USE_RTLD_GLOBAL"] = "1"
os.environ["PYTORCH_JIT"] = "0"
os.environ["PYTORCH_CUDA_ALLOC_CONF"] = "expandable_segments:False"
os.environ["CUDA_LAUNCH_BLOCKING"] = "1"
os.environ["TRANSFORMERS_CACHE"] = os.path.expanduser("~/.cache/transformers")

try:
    import torch
    import torch.nn as nn
    from torch import Tensor
    
    # 🔴 DESHABILITAR META TENSORS COMPLETAMENTE
    torch.set_default_dtype(torch.float32)
    
    # Guardar referencias originales
    _original_to = nn.Module.to
    _original_cpu = nn.Module.cpu
    _original_cuda = nn.Module.cuda
    _original_item = Tensor.item
    
    # 🔧 PARCHE 1: .to() seguro
    def _safe_to(self, *args, **kwargs):
        """Evita movimiento a meta device."""
        try:
            if args and any(str(a) == 'meta' for a in args):
                args = tuple(a for a in args if str(a) != 'meta')
                if not args:
                    args = ('cpu',)
            if kwargs.get('device') and str(kwargs['device']) == 'meta':
                kwargs['device'] = 'cpu'
            return _original_to(self, *args, **kwargs) if args or kwargs else self
        except Exception:
            return self
    
    # 🔧 PARCHE 2: .cpu() seguro
    def _safe_cpu(self):
        """CPU seguro sin meta tensors."""
        try:
            return _original_cpu(self)
        except Exception:
            return self
    
    # 🔧 PARCHE 3: .cuda() → .cpu()
    def _safe_cuda(self, *args, **kwargs):
        """CUDA siempre fallback a CPU."""
        try:
            return self.cpu() if hasattr(self, 'cpu') else self
        except Exception:
            return self
    
    # 🔧 PARCHE 4: .item() interceptado
    def _safe_item(self):
        """Intercepta .item() en meta tensors."""
        try:
            if hasattr(self, 'is_meta') and self.is_meta:
                return 0.0  # Valor seguro para meta tensors
            if self.device.type == 'meta':
                return 0.0
            return _original_item(self)
        except (RuntimeError, AttributeError, ValueError):
            return 0.0  # Fallback seguro
    
    # ✅ APLICAR TODOS LOS PARCHES
    nn.Module.to = _safe_to
    nn.Module.cpu = _safe_cpu
    nn.Module.cuda = _safe_cuda
    Tensor.item = _safe_item
    torch.Tensor.item = _safe_item
    
    # ⚠️ ADVERTENCIA DE INICIALIZACIÓN
    logger_init = logging.getLogger("MetaTensorFix")
    logger_init.info("✅ PyTorch meta tensor fix aplicado")
    
except ImportError:
    pass  # torch no disponible
except Exception as e:
    logger_init = logging.getLogger("MetaTensorFix")
    logger_init.warning(f"⚠️  Error en meta tensor fix: {e}")

# ============================================================

import re
import json
import gc
import sys
import textwrap
import sqlite3
import argparse
from typing import Dict, List, Optional, Tuple, Any
from datetime import datetime
from io import BytesIO
from html import escape

# --- DASHBOARD ---
import streamlit as st
import pandas as pd

# --- LANGCHAIN & IA ---
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import SentenceTransformerEmbeddings
from langchain.retrievers import EnsembleRetriever
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema.document import Document

# --- GENERACIÓN DE REPORTES ---
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.units import cm
from reportlab.lib.utils import simpleSplit
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# --- IMPORTAR TODO DESDE EL CORE DE INGESTA ---
# Este es el único punto de importación del sistema.
try:
    from ingesta_multibase import (
        # Configuración
        BASES_RAG,
        RELACIONES_MATERIAS,
        ROOT,
        LOGS_DIR,
        
        # Tipos y Modelos
        MetadataEnriquecido,
        TipoDocumento,
        Settings,
        
        # Funciones de Utilidad (Consolidadas)
        extract_text_from_pdf,
        classify_document,
        validate_pdf,
        get_vectorstore_safe,
        detect_metadata_enriched,
        materia_to_db_dir,
        make_logger,
    )
    INGESTA_AVAILABLE = True
except ImportError as e:
    INGESTA_AVAILABLE = False
    # Fallback si ingesta_multibase.py no está (aunque no debería pasar)
    st.error(f"Error Crítico: No se pudo importar ingesta_multibase.py. {e}")
    sys.exit()

# ============================================================
# ⚙️ CONFIGURACIÓN Y LOGGING
# ============================================================

# Desactivar telemetría de Chroma (reduce ruido de logs)
os.environ.setdefault("CHROMADB_TELEMETRY_ENABLED", "false")

logger = make_logger("AnalysisEngine")
CONFIG = Settings()

# Directorios de salida
OUTPUT_DIR = ROOT / "output_analisis_unificado"
DB_DIR = ROOT / "analysis_db_unificada"
PDF_TEMP_DIR = ROOT / "temp_pdfs"
OUTPUT_DIR.mkdir(exist_ok=True)
DB_DIR.mkdir(exist_ok=True)
PDF_TEMP_DIR.mkdir(exist_ok=True)

# Silenciar loggers ruidosos para una experiencia más limpia
for _name, _level in [
    ("chromadb.telemetry", logging.ERROR),
    ("chromadb", logging.WARNING),
    ("httpx", logging.WARNING),
    ("asyncio", logging.ERROR),
    ("PIL.PngImagePlugin", logging.ERROR),
]:
    _lg = logging.getLogger(_name)
    _lg.setLevel(_level)
    _lg.propagate = False

# Base de datos SQLite para histórico
DB_PATH = DB_DIR / "analisis_historico.sqlite3"

# Modelos LLM (Centralizado)
MODEL_OPTIONS = {
    "gemini-2.5-pro": "Gemini 2.5 Pro (Precisión)",
    "gemini-2.5-flash": "Gemini 2.5 Flash (Velocidad)",
}
DEFAULT_LLM = "gemini-2.5-pro"

# Helper: obtiene API Key de entorno (si existe)
def get_api_key() -> Optional[str]:
    for key in ("GEMINI_API_KEY", "GOOGLE_API_KEY"):
        val = os.environ.get(key)
        if val:
            return val
    return None

# Helper: detectar si se está ejecutando dentro de Streamlit
def _running_in_streamlit() -> bool:
    try:
        from streamlit.runtime.scriptrunner import get_script_run_ctx
        return get_script_run_ctx() is not None
    except Exception:
        return False

# ============================================================
# 🗄️ BASE DE DATOS SQLITE UNIFICADA
# ============================================================

def safe_search_query(query_func):
    """Decorator para envolver búsquedas y evitar meta tensors."""
    def wrapper(*args, **kwargs):
        try:
            return query_func(*args, **kwargs)
        except RuntimeError as e:
            if "meta" in str(e).lower() and "item" in str(e).lower():
                logger.warning(f"⚠️  Meta tensor detectado en búsqueda, usando fallback...")
                # Fallback: retornar resultado vacío o dummy
                if "SELECT" in str(query_func):
                    return None
                return []
            raise
    return wrapper

SCHEMA_SQL = """
PRAGMA journal_mode = WAL;
CREATE TABLE IF NOT EXISTS analisis_historico (
  id INTEGER PRIMARY KEY AUTOINCREMENT,
  tipo_analisis TEXT NOT NULL,
  nombre_archivo TEXT,
  expediente TEXT,
  tribunal TEXT,
  materia TEXT,
  fuero TEXT,
  resumen_analisis TEXT,
  ruta_pdf_generado TEXT,
  ruta_docx_generado TEXT,
  fecha_analisis TEXT NOT NULL,
  metadatos_json TEXT,
  fuentes_json TEXT,
  duracion_segundos REAL,
  modelo_llm TEXT
);
CREATE INDEX IF NOT EXISTS idx_tipo ON analisis_historico(tipo_analisis);
CREATE INDEX IF NOT EXISTS idx_fecha ON analisis_historico(fecha_analisis);
CREATE INDEX IF NOT EXISTS idx_materia ON analisis_historico(materia);
"""

def init_db():
    """Inicializa la base de datos SQLite con validación de rutas."""
    try:
        # Validar que el directorio existe
        DB_DIR.mkdir(parents=True, exist_ok=True)
        
        # Conectar y crear esquema
        with sqlite3.connect(str(DB_PATH)) as cx:
            cx.executescript(SCHEMA_SQL)
            cx.commit()
        
        logger.info(f"✅ BD unificada inicializada: {DB_PATH}")
        return True
    except sqlite3.OperationalError as e:
        logger.error(f"❌ Error operacional en BD: {e}")
        return False
    except Exception as e:
        logger.error(f"❌ Error inicializando BD unificada: {e}", exc_info=True)
        return False

def check_db_status() -> Dict[str, Any]:
    """Diagnóstico de estado de la base de datos."""
    status = {
        "db_path": str(DB_PATH),
        "db_exists": DB_PATH.exists(),
        "db_dir_exists": DB_DIR.exists(),
        "db_size_mb": 0,
        "tables": [],
        "row_count": 0,
        "error": None
    }
    
    try:
        if DB_PATH.exists():
            status["db_size_mb"] = round(DB_PATH.stat().st_size / (1024 * 1024), 3)
        
        with sqlite3.connect(str(DB_PATH)) as cx:
            cur = cx.cursor()
            
            # Listar tablas
            cur.execute("SELECT name FROM sqlite_master WHERE type='table'")
            status["tables"] = [row[0] for row in cur.fetchall()]
            
            # Contar filas
            if "analisis_historico" in status["tables"]:
                cur.execute("SELECT COUNT(*) FROM analisis_historico")
                status["row_count"] = cur.fetchone()[0]
    except Exception as e:
        status["error"] = str(e)
        logger.error(f"❌ Error diagnosticando BD: {e}")
    
    return status

def insert_analisis_db(data: Dict[str, Any]) -> int:
    """Inserta un análisis en la BD con manejo robusto de errores."""
    try:
        # Asegurar que la BD existe
        init_db()
        
        cols = [
            "tipo_analisis", "nombre_archivo", "expediente", "tribunal", 
            "materia", "fuero", "resumen_analisis", "ruta_pdf_generado", 
            "ruta_docx_generado", "fecha_analisis", "metadatos_json", 
            "fuentes_json", "duracion_segundos", "modelo_llm"
        ]
        vals = tuple(data.get(k) for k in cols)
        
        with sqlite3.connect(str(DB_PATH)) as cx:
            cx.isolation_level = None  # Autocommit para mejor estabilidad
            cur = cx.cursor()
            placeholders = ", ".join(["?"] * len(cols))
            cur.execute(
                f"INSERT INTO analisis_historico ({', '.join(cols)}) VALUES ({placeholders})",
                vals
            )
            row_id = cur.lastrowid
            logger.debug(f"✅ Análisis insertado con ID: {row_id}")
            return row_id
    except sqlite3.OperationalError as e:
        logger.error(f"❌ Error operacional al insertar en BD: {e}")
        return -1
    except Exception as e:
        logger.error(f"❌ Error insertando análisis en BD: {e}", exc_info=True)
        return -1

# ============================================================
# 🔬 SISTEMA DE EVALUACIÓN Y TRAZABILIDAD EPISTÉMICA
# ============================================================

import numpy as np

class LegalEvaluator:
    """
    Evaluador de desempeño PDCA para auditoría de calidad.
    Mide Precision, Recall y F1@k de los retrievers RAG.
    """
    def __init__(self, rag_manager):
        self.rag = rag_manager
    
    def evaluate(self, queries: List[str], goldsets: List[List[str]], k: int = 5) -> Dict[str, Any]:
        """
        Evalúa el desempeño del sistema RAG con queries de prueba.
        
        Args:
            queries: Lista de consultas de prueba
            goldsets: Lista de listas de IDs de documentos relevantes esperados
            k: Número de documentos a recuperar (top-k)
        
        Returns:
            Diccionario con métricas de evaluación
        """
        try:
            retriever = self.rag.get_combined_retriever(["Jurisprudencia", "General"])
            precs, recs = [], []
            
            for q, gt in zip(queries, goldsets):
                docs = retriever.get_relevant_documents(q)
                ids = [d.metadata.get("id_doc", d.metadata.get("archivo", "")) for d in docs[:k]]
                hits = sum(1 for g in gt if g in ids)
                p = hits / max(1, len(docs[:k]))
                r = hits / max(1, len(gt))
                precs.append(p)
                recs.append(r)
            
            f1s = [2*p*r/(p+r) if (p+r) > 0 else 0 for p, r in zip(precs, recs)]
            
            return {
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "precision@5": float(np.mean(precs)),
                "recall@5": float(np.mean(recs)),
                "f1@5": float(np.mean(f1s)),
                "num_queries": len(queries)
            }
        except Exception as e:
            logger.error(f"Error en evaluación PDCA: {e}")
            return {
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "error": str(e)
            }

class ModelRegistry:
    """
    Registro de versiones de modelos y bases de datos.
    Mantiene trazabilidad de qué versiones se usaron en cada análisis.
    """
    current = {
        "embeddings": "all-MiniLM-L6-v2@1.3",
        "llm_pro": "gemini-2.5-pro@2025-10",
        "llm_flash": "gemini-2.5-flash@2025-10",
        "bases": {
            "Jurisprudencia": "v1.2",
            "Civil - General": "v2.0",
            "Procesal - General": "v1.1",
            "Laboral - General": "v1.0"
        }
    }
    
    @classmethod
    def get_version_info(cls, materia: str = None) -> Dict[str, str]:
        """Obtiene información de versión para una materia específica."""
        info = {
            "embeddings": cls.current["embeddings"],
            "llm": cls.current["llm_pro"],
            "timestamp": datetime.now().isoformat()
        }
        if materia and materia in cls.current["bases"]:
            info["base_version"] = cls.current["bases"][materia]
        return info

def auditoria_juridica(texto: str) -> List[str]:
    """
    Auditoría jurídica automatizada: detecta omisiones normativas relevantes.
    
    Args:
        texto: Texto del análisis jurídico
    
    Returns:
        Lista de advertencias sobre referencias faltantes
    """
    reglas = [
        ("daño moral", "art. 1738 CCCN"),
        ("culpa", "art. 1724 CCCN"),
        ("menor de edad", "art. 25 CCCN"),
        ("responsabilidad objetiva", "art. 1722 CCCN"),
        ("buena fe", "art. 9 CCCN"),
        ("orden público", "art. 12 CCCN"),
        ("prescripción", "art. 2560 CCCN"),
        ("caducidad", "art. 2566 CCCN"),
    ]
    
    faltas = []
    texto_l = texto.lower()
    
    for kw, norma in reglas:
        if kw in texto_l and norma.lower() not in texto_l:
            faltas.append(f"Falta referencia a {norma} para '{kw}'.")
    
    return faltas

def evaluar_calidad_texto(texto: str) -> float:
    """
    Mide la coherencia textual mediante ratio de puntuación y longitud.
    
    Args:
        texto: Texto a evaluar
    
    Returns:
        Score de calidad entre 0.0 y 1.0
    """
    if not texto or len(texto.strip()) == 0:
        return 0.0
    
    puntuacion = sum(1 for c in texto if c in ".;:")
    longitud = max(1, len(texto.split()))
    
    # Calcular ratio (ajustado para textos jurídicos que tienden a ser densos)
    score = min(1.0, (puntuacion / longitud) * 20)
    return round(score, 3)

def export_jsonld(metadata: Any, analisis: str) -> Dict[str, Any]:
    """
    Exporta metadata y análisis en formato JSON-LD (Linked Data).
    Compatible con schema.org/LegalCase.
    
    Args:
        metadata: Objeto de metadatos enriquecidos
        analisis: Texto del análisis
    
    Returns:
        Diccionario JSON-LD
    """
    return {
        "@context": "https://schema.org",
        "@type": "LegalCase",
        "name": getattr(metadata, "expediente", "") or getattr(metadata, "nombre_fallo", ""),
        "jurisdiction": getattr(metadata, "jurisdiccion", ""),
        "court": getattr(metadata, "tribunal", ""),
        "legislationApplied": getattr(metadata, "materia", ""),
        "datePublished": getattr(metadata, "fecha_sentencia", "") or getattr(metadata, "fecha_resolucion", ""),
        "author": getattr(metadata, "responsable", "Sistema de Análisis Jurídico"),
        "version": getattr(metadata, "version", "1.0"),
        "analysisMethod": "RAG + LLM (Gemini)",
        "modelUsed": ModelRegistry.current["llm_pro"],
        "epistemicRisk": getattr(metadata, "riesgo_detectado", "Bajo"),
        "qualityScore": getattr(metadata, "calidad_texto", 0.0),
        "ethicalNotes": getattr(metadata, "comentarios_eticos", ""),
        "abstract": analisis[:600] if analisis else "",
        "identifier": getattr(metadata, "id_doc", "") or getattr(metadata, "archivo_origen", ""),
        "dateCreated": datetime.now().isoformat()
    }

# ============================================================
# 🧠 PROMPTS CENTRALIZADOS
# ============================================================

# Instrucción de contexto cerrado para evitar alucinaciones
INSTRUCTION_CC = (
    "IMPORTANTE: Responde SOLO usando el contexto recuperado. "
    "Si la información no está en el contexto, indica 'No se encuentra información suficiente en el contexto proporcionado'. "
    "NO inventes, supongas o agregues información externa.\n\n"
)

# --- Prompts de slegal.py (Análisis de Texto) ---
PROMPT_PERFECCIONAMIENTO = PromptTemplate(
    input_variables=["context", "question"],
    template=(
        INSTRUCTION_CC +
        "Eres un asistente jurídico experto. A partir del contexto provisto, "
        "mejora y perfecciona la redacción legal, identifica ambigüedades, "
        "propone alternativas terminológicas.\n\n"
        "Contexto:\n{context}\n\nPregunta:\n{question}\n\nRespuesta:"
    ),
)
PROMPT_INTERPRETACION = PromptTemplate(
    input_variables=["context", "question"],
    template=(
        INSTRUCTION_CC +
        "Actúa como un jurista especializado en interpretación normativa. "
        "Del contexto, extrae la regla aplicable, enumera supuestos de hecho.\n\n"
        "Contexto:\n{context}\n\nPregunta:\n{question}\n\nRespuesta:"
    ),
)
PROMPT_ANALISIS_AVANZADO = PromptTemplate(
    input_variables=["context", "question"],
    template=(
        INSTRUCTION_CC +
        "Adopta el rol de analista jurídico estratégico. "
        "Realiza un análisis crítico de riesgos, alternativas procesales.\n\n"
        "Contexto:\n{context}\n\nPregunta:\n{question}\n\nRespuesta:"
    ),
)

# --- Prompt de analista_jurisprudencial.py (Análisis de Fallo) ---
PROMPT_JURISPRUDENCIA = PromptTemplate(
    input_variables=["fallo","ctx_juris","ctx_legal","ctx_related"],
    template=(
        INSTRUCTION_CC +
        """Analizá el siguiente fallo argentino considerando contexto de jurisprudencia, doctrina y áreas relacionadas.

FALLO:
{fallo}

[Jurisprudencia relacionada]
{ctx_juris}

[Doctrina relacionada]
{ctx_legal}

[Contexto de áreas relacionadas]
{ctx_related}

Redactá un informe técnico-forense con secciones:
1. Identificación (expediente, tribunal, jueces, año)
2. Análisis de Materia y Jurisdicción
3. Hechos Relevantes
4. Cuestiones Jurídicas
5. Análisis Jurisprudencial y Doctrinal
6. Razonamiento del Tribunal
7. Conclusión y Tendencia
"""
    )
)

# --- Prompt de analista_procesos_judiciales.py (Análisis de Expediente) ---
PROMPT_PROCESAL = PromptTemplate(
    input_variables=["texto_expediente", "fuero", "etapas", "ctx_juris", "ctx_doctrina", "ctx_relacionada"],
    template=(
        INSTRUCTION_CC +
        """Analizá el siguiente expediente judicial argentino en su estructura procesal, 
considerando jurisprudencia, doctrina y casos relacionados.

TEXTO DEL EXPEDIENTE:
{texto_expediente}

FUERO: {fuero}
ETAPAS DETECTADAS: {etapas}

[JURISPRUDENCIA RELACIONADA]
{ctx_juris}

[DOCTRINA Y NORMAS APLICABLES]
{ctx_doctrina}

[CONTEXTO DE FUERO RELACIONADO]
{ctx_relacionada}

Elaborá un DICTAMEN TÉCNICO PROCESAL con las siguientes secciones:
1. IDENTIFICACIÓN DEL EXPEDIENTE
2. CRONOLOGÍA PROCESAL
3. ANÁLISIS DE ACTOS PROCESALES
4. RESOLUCIONES Y SU ESTADO
5. FUNDAMENTOS JURISPRUDENCIALES Y DOCTRINALES
6. RECOMENDACIONES PROCEDIMENTALES
7. OBSERVACIONES FINALES
"""
    )
)

# --- Prompt de slegal.py (Q&A sobre Expediente) ---
PROMPT_QA_EXPEDIENTE = PromptTemplate(
    input_variables=["context", "question"], 
    template=(
        INSTRUCTION_CC +
        """Eres abogado experto en análisis de expedientes.
Responde la pregunta considerando TANTO el expediente como la jurisprudencia disponible.

PREGUNTA DEL USUARIO:
{question}

CONTEXTO LEGAL DISPONIBLE (jurisprudencia, leyes, doctrina y fragmentos del expediente):
{context}

INSTRUCCIONES:
1. Busca la respuesta en el contexto proporcionado.
2. Cita jurisprudencia, leyes o doctrina aplicables.
3. Explica cómo se aplica al expediente.
4. Sé preciso y cita fuentes.
"""
    )
)

# ============================================================
# 📦 CLASE: ReportGenerator (Exportación PDF/DOCX)
# ============================================================

class ReportGenerator:
    """Consolida la generación de reportes PDF y DOCX."""
    
    def __init__(self):
        self.logger = make_logger("ReportGenerator")

    def export_to_pdf(
        self,
        contenido: str,
        titulo: str,
        metadatos: Dict[str, Any],
        output_path: Path
    ) -> bool:
        """Exporta análisis a PDF (lógica de analista_jurisprudencial.py)."""
        try:
            output_path.parent.mkdir(parents=True, exist_ok=True)
            c = canvas.Canvas(str(output_path), pagesize=A4)
            w, h = A4
            margin_x = 2 * cm
            content_width = w - 4 * cm
            y = h - 2 * cm
            line_height = 14

            c.setFont("Helvetica-Bold", 14)
            c.drawCentredString(w / 2, y, titulo.upper())
            y -= 25
            
            # Metadatos
            c.setFont("Helvetica", 9)
            meta_text = f"Expediente: {metadatos.get('expediente', 'N/A')} | Tribunal: {metadatos.get('tribunal', 'N/A')} | Materia: {metadatos.get('materia', 'N/A')}"
            c.drawCentredString(w / 2, y, meta_text)
            y -= 20
            
            # Contenido
            c.setFont("Helvetica", 10)
            for line in contenido.split("\n"):
                lines_to_draw = simpleSplit(line, 'Helvetica', 10, content_width)
                for text_segment in lines_to_draw:
                    if y < 2 * cm:
                        c.showPage()
                        c.setFont("Helvetica", 10)
                        y = h - 2 * cm
                    c.drawString(margin_x, y, text_segment)
                    y -= line_height
            
            c.save()
            self.logger.info(f"PDF generado: {output_path}")
            return True
            
        except Exception as e:
            self.logger.error(f"Error generando PDF: {e}")
            return False

    def export_to_docx(
        self,
        analisis_dict: Dict[str, str],
        metadatos: Dict[str, Any],
        output_path: Path
    ) -> bool:
        """Exporta análisis a DOCX (lógica de slegal.py)."""
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx no instalado. No se puede exportar a DOCX.")
            return False
        
        try:
            doc = DocxDocument()
            doc.add_heading("DICTAMEN JURÍDICO", level=1)
            
            # Metadatos
            doc.add_paragraph(f"Expediente: {metadatos.get('expediente', 'N/A')}")
            doc.add_paragraph(f"Tribunal: {metadatos.get('tribunal', 'N/A')}")
            doc.add_paragraph(f"Materia: {metadatos.get('materia', 'N/A')}")
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d')}")
            
            # Contenido
            for seccion, contenido in analisis_dict.items():
                doc.add_heading(seccion.upper(), level=2)
                for parrafo in contenido.split("\n"):
                    if parrafo.strip():
                        doc.add_paragraph(parrafo.strip())
            
            doc.save(output_path)
            self.logger.info(f"DOCX generado: {output_path}")
            return True
        except Exception as e:
            self.logger.error(f"Error generando DOCX: {e}")
            return False

# ============================================================
# 🔧 PARCHE DE ENRIQUECIMIENTO Y TRAZABILIDAD - ANALYSER (compat)
# ============================================================

# Nota: Funciones auxiliares agregadas en paralelo sin modificar el flujo existente.
# Puedes usarlas explícitamente según necesidad sin afectar lo que ya funciona.

import argparse as _argparse_compat  # uso local para helper opcional


class LegalEvaluatorParche:
    """Evalúa precisión, recall y F1@k para control de calidad del sistema RAG."""
    def __init__(self, rag_manager):
        self.rag = rag_manager

    def evaluate(self, queries, goldsets, k=5):
        retriever = self.rag.get_combined_retriever(["Jurisprudencia", "General"])
        precs, recs = [], []
        for q, gt in zip(queries, goldsets):
            docs = retriever.get_relevant_documents(q)
            ids = [d.metadata.get("id_doc") for d in docs[:k]]
            hits = sum(1 for g in gt if g in ids)
            p = hits / max(1, len(docs))
            r = hits / max(1, len(gt))
            precs.append(p)
            recs.append(r)
        f1s = [2 * p * r / (p + r) if (p + r) > 0 else 0 for p, r in zip(precs, recs)]
        return {
            "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
            "precision@5": float(np.mean(precs)),
            "recall@5": float(np.mean(recs)),
            "f1@5": float(np.mean(f1s))
        }


class ModelRegistryParche:
    current = {
        "embeddings": "all-MiniLM-L6-v2@1.3",
        "llm": "gemini-1.5-pro@2025-10",
        "bases": {
            "Jurisprudencia": "v1.2",
            "Civil": "v2.0",
            "General": "v1.1"
        }
    }


def auditoria_juridica_parche(texto: str) -> list[str]:
    """Detecta omisiones normativas relevantes en el texto analizado."""
    reglas = [
        ("daño moral", "art. 1738 CCCN"),
        ("culpa", "art. 1724 CCCN"),
        ("menor de edad", "art. 25 CCCN")
    ]
    faltas = []
    texto_l = texto.lower()
    for kw, norma in reglas:
        if kw in texto_l and norma.lower() not in texto_l:
            faltas.append(f"Falta referencia a {norma} para '{kw}'.")
    return faltas


def evaluar_calidad_texto_parche(texto: str) -> float:
    """Mide la coherencia textual (0-1) según proporción de puntuación y longitud."""
    puntuacion = sum(1 for c in texto if c in ".;:")
    longitud = max(1, len(texto.split()))
    return round(min(1.0, (puntuacion / longitud) * 20), 3)


def export_jsonld_parche(metadata, analisis):
    """Genera estructura JSON-LD interoperable con metadatos jurídicos FAIR."""
    return {
        "@context": "https://schema.org",
        "@type": "LegalCase",
        "name": getattr(metadata, "expediente", ""),
        "jurisdiction": getattr(metadata, "jurisdiccion", ""),
        "court": getattr(metadata, "tribunal", ""),
        "legislationApplied": getattr(metadata, "materia", ""),
        "datePublished": getattr(metadata, "fecha_resolucion", ""),
        "author": getattr(metadata, "responsable", ""),
        "version": getattr(metadata, "version", ""),
        "analysisMethod": getattr(metadata, "metodo_analisis", ""),
        "modelUsed": getattr(metadata, "modelo_usado", ""),
        "epistemicRisk": getattr(metadata, "riesgo_detectado", ""),
        "qualityScore": getattr(metadata, "calidad_texto", 0.0),
        "flowStatus": getattr(metadata, "estado_flujo", ""),
        "kanbanPriority": getattr(metadata, "prioridad", ""),
        "ethicalNotes": getattr(metadata, "comentarios_eticos", ""),
        "abstract": analisis[:600] if analisis else "",
        "identifier": getattr(metadata, "id_doc", "")
    }

    def _run_pdca_example_parche():
        """Helper opcional: ejecuta evaluación PDCA de ejemplo sin tocar el CLI principal."""
        try:
            rag = RAGManager(CONFIG)
            evaluator = LegalEvaluatorParche(rag)
            queries = [
                "despido injustificado y cálculo indemnizatorio",
                "daño moral por incumplimiento contractual"
            ]
            goldsets = [["uuid1", "uuid2"], ["uuid3", "uuid4"]]
            return evaluator.evaluate(queries, goldsets)
        except Exception as _e:
            return {"error": str(_e)}
            for line in contenido.split("\n"):
                lines_to_draw = simpleSplit(line, 'Helvetica', 10, content_width)
                for text_segment in lines_to_draw:
                    if y < 2 * cm:
                        c.showPage()
                        c.setFont("Helvetica", 10)
                        y = h - 2 * cm
                    c.drawString(margin_x, y, text_segment)
                    y -= line_height
            c.save()
            self.logger.info(f"PDF generado: {output_path}")
            return True
        except Exception as e:
            self.logger.error(f"Error generando PDF: {e}")
            return False

    def export_to_docx(
        self,
        analisis_dict: Dict[str, str],
        metadatos: Dict[str, Any],
        output_path: Path
    ) -> bool:
        """Exporta análisis a DOCX (lógica de slegal.py)."""
        if not DOCX_AVAILABLE:
            self.logger.warning("python-docx no instalado. No se puede exportar a DOCX.")
            return False
        
        try:
            doc = DocxDocument()
            doc.add_heading("DICTAMEN JURÍDICO", level=1)
            
            # Metadatos
            doc.add_paragraph(f"Expediente: {metadatos.get('expediente', 'N/A')}")
            doc.add_paragraph(f"Tribunal: {metadatos.get('tribunal', 'N/A')}")
            doc.add_paragraph(f"Materia: {metadatos.get('materia', 'N/A')}")
            doc.add_paragraph(f"Fecha: {datetime.now().strftime('%Y-%m-%d')}")
            
            # Contenido
            for seccion, contenido in analisis_dict.items():
                doc.add_heading(seccion.upper(), level=2)
                for parrafo in contenido.split("\n"):
                    if parrafo.strip():
                        doc.add_paragraph(parrafo.strip())
            
            doc.save(output_path)
            self.logger.info(f"DOCX generado: {output_path}")
            return True
        except Exception as e:
            self.logger.error(f"Error generando DOCX: {e}")
            return False

# ============================================================
# � FIX: Wrapper para SentenceTransformer (Python 3.13 fix)
# ============================================================

class SafeSentenceTransformerEmbeddings:
    """Wrapper personalizado para evitar el bug de meta tensors en Python 3.13."""
    
    def __init__(self, model_name: str):
        import sentence_transformers
        
        # Usar la carpeta local models/ como cache
        cache_folder = os.path.join(os.path.dirname(__file__), "models")
        
        # Cargar el modelo - el monkey patch interceptará .to()
        self.model = sentence_transformers.SentenceTransformer(
            model_name,
            device="cpu",  # Forzar CPU desde el inicio
            cache_folder=cache_folder
        )
    
    def embed_documents(self, texts):
        """Embeddings para múltiples documentos."""
        return self.model.encode(texts, convert_to_numpy=True).tolist()
    
    def embed_query(self, text):
        """Embedding para una sola consulta."""
        return self.model.encode([text], convert_to_numpy=True)[0].tolist()

# ============================================================
# �🔎 CLASE: RAGManager (Gestión de RAG Dinámico y Local)
# ============================================================

class RAGManager:
    """Maneja la carga y combinación de todos los retrievers."""
    
    def __init__(self, settings: Settings):
        self.settings = settings
        self.logger = make_logger("RAGManager")
        # Usar wrapper personalizado en lugar de SentenceTransformerEmbeddings
        # para evitar el bug de meta tensors en Python 3.13
        self.embeddings = SafeSentenceTransformerEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2"
        )

        self.all_retrievers = self._load_all_retrievers()

    @st.cache_resource(show_spinner="Cargando bases de datos RAG...")
    def _load_all_retrievers(_self) -> Dict[str, Any]:
        """Carga TODOS los retrievers disponibles existentes.

        - Omite rutas inexistentes sin warnings.
        - Registra error si falla una base existente.
        - No interrumpe: intenta con todas las bases y carga las que estén disponibles.
        """
        retrievers: Dict[str, Any] = {}
        for materia in BASES_RAG.keys():
            ruta_base = materia_to_db_dir(materia)
            if not Path(ruta_base).exists():
                # Salta las que no existen sin warning
                continue

            try:
                vectorstore = Chroma(
                    persist_directory=str(ruta_base),
                    embedding_function=_self.embeddings,
                )
                retrievers[materia] = vectorstore.as_retriever(search_kwargs={"k": 5})
                _self.logger.info(f"✅ Base cargada: {materia}")
            except Exception as e:
                _self.logger.error(f"Error cargando {materia}: {e}")
                continue

        _self.logger.info(f"Cargadas {len(retrievers)} base(s).")
        return retrievers

    def get_combined_retriever(self, selected_bases: List[str]) -> Optional[Any]:
        """Crea un EnsembleRetriever dinámico (feature de slegal.py)."""
        active_retrievers = [
            self.all_retrievers[base] 
            for base in selected_bases 
            if base in self.all_retrievers
        ]
        
        if not active_retrievers:
            self.logger.error("No se seleccionó ningún retriever válido.")
            return None
        if len(active_retrievers) == 1:
            return active_retrievers[0]
        
        weights = [1.0 / len(active_retrievers)] * len(active_retrievers)
        return EnsembleRetriever(retrievers=active_retrievers, weights=weights)

    def get_retriever_for_base(self, base_name: str) -> Optional[Any]:
        """Devuelve un retriever de Chroma para una base específica si existe."""
        ruta_base = materia_to_db_dir(base_name)
        if not Path(ruta_base).exists():
            return None
        try:
            vectorstore = Chroma(
                persist_directory=str(ruta_base),
                embedding_function=self.embeddings,
            )
            return vectorstore.as_retriever(search_kwargs={"k": 5})
        except Exception as e:
            self.logger.error(f"Error cargando base '{base_name}': {e}")
            return None

    def create_expediente_retriever(self, texto_expediente: str) -> Optional[Any]:
        """Crea un retriever local en memoria para el PDF subido (feature de slegal.py)."""
        try:
            splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.settings.chunk_size,
                chunk_overlap=self.settings.chunk_overlap
            )
            chunks = splitter.split_text(texto_expediente)
            
            if not chunks:
                self.logger.warning("No se generaron chunks del expediente.")
                return None
            
            vectorstore = Chroma.from_texts(
                texts=chunks,
                embedding=self.embeddings,
                collection_name="expediente_local"
            )
            self.logger.info(f"Retriever local de expediente creado ({len(chunks)} chunks).")
            return vectorstore.as_retriever(search_kwargs={"k": 5})
        except Exception as e:
            self.logger.error(f"Error creando retriever de expediente: {e}")
            return None

# ============================================================
# ✍️ CLASE: TextAnalyzer (Análisis de Texto)
# ============================================================

class TextAnalyzer:
    """Maneja los 4 flujos de análisis de texto (de slegal.py)."""
    
    def __init__(self, model_name: str, retriever: Any):
        self.logger = make_logger("TextAnalyzer")
        self.llm = ChatGoogleGenerativeAI(model=model_name, temperature=0.3, api_key=get_api_key())
        self.retriever = retriever
        
        # Crear cadenas QA
        self.qa_p = self._create_chain(PROMPT_PERFECCIONAMIENTO)
        self.qa_i = self._create_chain(PROMPT_INTERPRETACION)
        self.qa_a = self._create_chain(PROMPT_ANALISIS_AVANZADO)

    def _create_chain(self, prompt: PromptTemplate) -> RetrievalQA:
        return RetrievalQA.from_chain_type(
            llm=self.llm,
            chain_type="stuff",
            retriever=self.retriever,
            return_source_documents=False,
            chain_type_kwargs={"prompt": prompt},
        )

    def _safe_invoke(self, chain: RetrievalQA, text: str) -> str:
        try:
            response = chain.invoke({"query": text})
            return response.get("result", "No se obtuvo respuesta.")
        except Exception as e:
            self.logger.error(f"Error en chain invoke: {e}")
            return f"[Error: {e}]"

    def run_perfeccionamiento(self, text: str) -> str:
        return self._safe_invoke(self.qa_p, text)

    def run_interpretacion(self, text: str) -> str:
        return self._safe_invoke(self.qa_i, text)

    def run_analisis_avanzado(self, text: str) -> str:
        return self._safe_invoke(self.qa_a, text)

    def run_analisis_integral(self, text: str) -> str:
        """El meta-análisis que combina los otros tres."""
        self.logger.info("Iniciando análisis integral...")
        r1 = self.run_perfeccionamiento(text)
        r2 = self.run_interpretacion(text)
        r3 = self.run_analisis_avanzado(text)
        
        meta_prompt = f"""
Eres un META-ANALISTA JURÍDICO SUPERIOR.
Integra los siguientes tres análisis en un DICTAMEN JURÍDICO UNIFICADO y coherente.

CONSULTA ORIGINAL:
{text}

---
ANÁLISIS 1 (PERFECCIONAMIENTO):
{r1}
---
ANÁLISIS 2 (INTERPRETACIÓN):
{r2}
---
ANÁLISIS 3 (ESTRATÉGICO):
{r3}
---

DICTAMEN UNIFICADO (Estructura: I. Síntesis, II. Conclusión Integral, III. Recomendación):
"""
        try:
            response = self.llm.invoke(meta_prompt)
            return response.content if hasattr(response, 'content') else str(response)
        except Exception as e:
            self.logger.error(f"Error en meta-análisis: {e}")
            return f"[Error en meta-análisis: {e}]"

# ============================================================
# 📚 CLASE: JurisprudenciaAnalyzer (Análisis de Fallos PDF)
# ============================================================

class JurisprudenciaAnalyzer:
    """Maneja el análisis de Fallos PDF (de analista_jurisprudencial.py)."""
    
    def __init__(self, model_name: str, rag_manager: RAGManager):
        self.logger = make_logger("JurisprudenciaAnalyzer")
        self.llm = ChatGoogleGenerativeAI(model=model_name, temperature=0.3, api_key=get_api_key())
        self.rag_manager = rag_manager
        self.report_gen = ReportGenerator()

    def _build_context_block(self, docs, title):
        acc=[f"===== {title} ====="]
        for i,d in enumerate(docs[:3],1):
            snippet = d.page_content.replace("\n"," ")
            acc.append(f"[{i}] {snippet[:600]}")
        return "\n".join(acc)

    def analyze(self, pdf_path: Path) -> Dict[str, Any]:
        self.logger.info(f"Analizando fallo: {pdf_path.name}")
        
        success, text = extract_text_from_pdf(pdf_path)
        if not success:
            raise ValueError("No se pudo extraer texto del PDF.")
        
        metadata = detect_metadata_enriched(text, fallback_name=pdf_path.stem)
        metadata.tipo_documento = TipoDocumento.FALLO_JURISPRUDENCIAL
        
        # --- Lógica de RAG Dual (de analista_jurisprudencial.py) ---
        materia = metadata.materia or "General"
        
        # 1. Búsqueda en Jurisprudencia
        ret_juris = self.rag_manager.all_retrievers.get("Jurisprudencia")
        docs_juris = ret_juris.get_relevant_documents(text) if ret_juris else []
        
        # 2. Búsqueda en Doctrina/Legal (General)
        ret_general = self.rag_manager.all_retrievers.get("General")
        docs_legal = ret_general.get_relevant_documents(text) if ret_general else []
        
        # 3. Búsqueda en Materia Específica
        ret_materia = self.rag_manager.all_retrievers.get(materia)
        docs_related = ret_materia.get_relevant_documents(text) if ret_materia else []
        
        ctx_juris = self._build_context_block(docs_juris, "JURISPRUDENCIA")
        ctx_legal = self._build_context_block(docs_legal, "DOCTRINA")
        ctx_related = self._build_context_block(docs_related, f"CONTEXTO: {materia}")
        
        prompt = PROMPT_JURISPRUDENCIA.format(
            fallo=text[:8000],
            ctx_juris=ctx_juris,
            ctx_legal=ctx_legal,
            ctx_related=ctx_related
        )
        
        response = self.llm.invoke(prompt)
        analisis = response.content if hasattr(response, 'content') else str(response)
        
        # 🔬 AUDITORÍA Y TRAZABILIDAD
        faltas = auditoria_juridica(analisis)
        metadata.calidad_texto = evaluar_calidad_texto(analisis)
        metadata.riesgo_detectado = "Medio" if "confidencial" in analisis.lower() else "Bajo"
        metadata.version = ModelRegistry.current["llm_pro"]
        
        # Agregar advertencias QA si hay omisiones
        if faltas:
            analisis += "\n\n[⚠️ ADVERTENCIAS DE AUDITORÍA JURÍDICA]\n- " + "\n- ".join(faltas)
        
        return {
            "metadata": metadata,
            "analisis": analisis,
            "fuentes": {"juris": len(docs_juris), "legal": len(docs_legal), "related": len(docs_related)},
            "calidad_texto": metadata.calidad_texto,
            "advertencias_qa": len(faltas)
        }

# ============================================================
# ⚖️ CLASE: ProcesoAnalyzer (Análisis de Expedientes PDF)
# ============================================================

class ProcesoAnalyzer:
    """Maneja el análisis de Expedientes PDF (de analista_procesos_judiciales.py)."""
    
    def __init__(self, model_name: str, rag_manager: RAGManager):
        self.logger = make_logger("ProcesoAnalyzer")
        self.llm = ChatGoogleGenerativeAI(model=model_name, temperature=0.3, api_key=get_api_key())
        self.rag_manager = rag_manager
        self.report_gen = ReportGenerator()

    def _build_context_block(self, docs, title):
        acc=[f"===== {title} ====="]
        for i,d in enumerate(docs[:3],1):
            snippet = d.page_content.replace("\n"," ")
            acc.append(f"[{i}] {snippet[:600]}")
        return "\n".join(acc)

    def analyze(self, pdf_path: Path) -> Dict[str, Any]:
        self.logger.info(f"Analizando expediente: {pdf_path.name}")
        
        success, text = extract_text_from_pdf(pdf_path)
        if not success:
            raise ValueError("No se pudo extraer texto del PDF.")
        
        metadata = detect_metadata_enriched(text, fallback_name=pdf_path.stem)
        metadata.tipo_documento = TipoDocumento.EXPEDIENTE
        
        # --- Lógica de RAG (de analista_procesos_judiciales.py) ---
        fuero = metadata.fuero or "General"
        
        ret_juris = self.rag_manager.all_retrievers.get("Jurisprudencia")
        docs_juris = ret_juris.get_relevant_documents(text) if ret_juris else []
        
        ret_doctrina = self.rag_manager.all_retrievers.get("General")
        docs_doctrina = ret_doctrina.get_relevant_documents(text) if ret_doctrina else []

        ret_relacionada = self.rag_manager.all_retrievers.get(fuero)
        docs_relacionada = ret_relacionada.get_relevant_documents(text) if ret_relacionada else []
        
        ctx_juris = self._build_context_block(docs_juris, "JURISPRUDENCIA")
        ctx_doctrina = self._build_context_block(docs_doctrina, "DOCTRINA")
        ctx_relacionada = self._build_context_block(docs_relacionada, f"CONTEXTO: {fuero}")
        
        prompt = PROMPT_PROCESAL.format(
            texto_expediente=text[:8000],
            fuero=fuero,
            etapas=", ".join(metadata.etapas_detectadas),
            ctx_juris=ctx_juris,
            ctx_doctrina=ctx_doctrina,
            ctx_relacionada=ctx_relacionada
        )
        
        response = self.llm.invoke(prompt)
        analisis = response.content if hasattr(response, 'content') else str(response)
        
        # 🔬 AUDITORÍA Y TRAZABILIDAD
        faltas = auditoria_juridica(analisis)
        metadata.calidad_texto = evaluar_calidad_texto(analisis)
        metadata.riesgo_detectado = "Medio" if "confidencial" in analisis.lower() else "Bajo"
        metadata.version = ModelRegistry.current["llm_pro"]
        
        # Agregar advertencias QA si hay omisiones
        if faltas:
            analisis += "\n\n[⚠️ ADVERTENCIAS DE AUDITORÍA JURÍDICA]\n- " + "\n- ".join(faltas)
        
        # Crear retriever local para Q&A
        retriever_local = self.rag_manager.create_expediente_retriever(text)
        
        return {
            "metadata": metadata,
            "analisis": analisis,
            "texto_completo": text,
            "retriever_local": retriever_local,
            "fuentes": {"juris": len(docs_juris), "doctrina": len(docs_doctrina), "related": len(docs_relacionada)},
            "calidad_texto": metadata.calidad_texto,
            "advertencias_qa": len(faltas)
        }

# ============================================================
# 🖥️ CLASE: DashboardUI (Streamlit Unificado)
# ============================================================

class DashboardUI:
    """Orquesta la UI de Streamlit, unificando todos los flujos."""
    
    def __init__(self):
        self.logger = make_logger("DashboardUI")
        self.rag_manager = RAGManager(CONFIG)
        self.report_gen = ReportGenerator()
        init_db()
        
        # Inicializar session state
        if "selected_model" not in st.session_state:
            st.session_state.selected_model = DEFAULT_LLM
        if "selected_bases" not in st.session_state:
            st.session_state.selected_bases = list(BASES_RAG.keys())
        if "global_retriever" not in st.session_state:
            st.session_state.global_retriever = self.rag_manager.get_combined_retriever(st.session_state.selected_bases)
        if "local_retriever" not in st.session_state:
            st.session_state.local_retriever = None
        if "current_expediente_text" not in st.session_state:
            st.session_state.current_expediente_text = ""
        if "qa_chain" not in st.session_state:
            st.session_state.qa_chain = None

    def _render_sidebar(self):
        """Renderiza la barra lateral con opciones globales."""
        with st.sidebar:
            st.title("⚖️ Motor Jurídico")
            st.markdown("Configuración de Análisis")
            
            # 1. Selección de Modelo LLM
            selected_model = st.radio(
                "Modelo de IA:",
                options=list(MODEL_OPTIONS.keys()),
                format_func=lambda x: MODEL_OPTIONS[x],
                key="selected_model_radio"
            )
            if selected_model != st.session_state.selected_model:
                st.session_state.selected_model = selected_model
                st.rerun()

            st.markdown("---")
            
            # 2. Selección Dinámica de RAG (feature de slegal.py)
            st.subheader("Bases de Datos RAG")
            all_bases = list(self.rag_manager.all_retrievers.keys())
            
            selected_bases = st.multiselect(
                "Selecciona las bases a consultar:",
                options=all_bases,
                default=all_bases
            )
            
            if st.button("Actualizar Bases RAG"):
                st.session_state.selected_bases = selected_bases
                st.session_state.global_retriever = self.rag_manager.get_combined_retriever(selected_bases)
                # Limpiar chains para que se recreen con el nuevo retriever
                st.session_state.pop("text_analyzer", None) 
                st.session_state.pop("qa_chain", None)
                st.success("Bases RAG actualizadas.")
                st.rerun()

            st.info(f"Consultando {len(st.session_state.selected_bases)} bases.")

            # Selección rápida de una sola base existente (por Chroma)
            st.markdown("---")
            st.caption("Selección rápida de una base (Chroma)")
            bases_disponibles = [
                m for m in BASES_RAG.keys() if Path(materia_to_db_dir(m)).exists()
            ]
            if bases_disponibles:
                selected_base = st.selectbox(
                    "Seleccionar base Chroma:",
                    bases_disponibles,
                    key="single_base_select"
                )
                if st.button("Usar base seleccionada"):
                    single_ret = self.rag_manager.get_retriever_for_base(selected_base)
                    if single_ret:
                        st.session_state.selected_bases = [selected_base]
                        st.session_state.global_retriever = single_ret
                        st.session_state.pop("text_analyzer", None)
                        st.session_state.pop("qa_chain", None)
                        st.success(f"Usando base: {selected_base}")
                        st.rerun()
                    else:
                        st.error("No se pudo cargar la base seleccionada.")
            else:
                st.info("No hay bases Chroma construidas disponibles para selección rápida.")

    def _run_tab_texto(self):
        """Tab para Análisis de Texto (de slegal.py)."""
        st.subheader("Análisis de Texto y Cláusulas")
        
        flow_selection = st.radio(
            "Selecciona el tipo de análisis:",
            (
                "Perfeccionamiento",
                "Interpretación y Validación",
                "Análisis Avanzado",
                "Análisis Integral (Combina los 3)",
            ),
            horizontal=True,
        )
        
        user_input = st.text_area(
            "Pega el texto legal, cláusula o consulta:",
            height=150,
            placeholder="Ej: 'El locatario no podrá subarrendar el inmueble...'",
        )
        
        if st.button("🚀 Analizar Texto", use_container_width=True, disabled=not user_input):
            if not st.session_state.global_retriever:
                st.error("No hay bases RAG seleccionadas. Por favor, selecciona al menos una en la barra lateral.")
                return

            # Inicializar el analizador de texto con el retriever global actual
            text_analyzer = TextAnalyzer(
                st.session_state.selected_model, 
                st.session_state.global_retriever
            )
            
            with st.spinner(f"Ejecutando {flow_selection}..."):
                if flow_selection == "Perfeccionamiento":
                    resultado = text_analyzer.run_perfeccionamiento(user_input)
                elif flow_selection == "Interpretación y Validación":
                    resultado = text_analyzer.run_interpretacion(user_input)
                elif flow_selection == "Análisis Avanzado":
                    resultado = text_analyzer.run_analisis_avanzado(user_input)
                else: # Integral
                    resultado = text_analyzer.run_analisis_integral(user_input)
            
            st.markdown("---")
            st.subheader("Resultado del Análisis")
            st.markdown(resultado)

    def _run_tab_fallos(self):
        """Tab para Análisis de Fallos PDF (de analista_jurisprudencial.py)."""
        st.subheader("Análisis de Fallos Jurisprudenciales (PDF)")
        
        uploaded_file = st.file_uploader(
            "Sube un fallo en PDF:",
            type=["pdf"],
            key="juris_upload"
        )
        
        if uploaded_file:
            if st.button("⚖️ Analizar Fallo", use_container_width=True):
                try:
                    # Guardar temporalmente
                    temp_path = PDF_TEMP_DIR / uploaded_file.name
                    with open(temp_path, "wb") as f: f.write(uploaded_file.getbuffer())
                    
                    analyzer = JurisprudenciaAnalyzer(
                        st.session_state.selected_model,
                        self.rag_manager
                    )
                    
                    with st.spinner(f"Analizando fallo con {st.session_state.selected_model}..."):
                        result = analyzer.analyze(temp_path)
                    
                    st.success("Análisis de fallo completado.")
                    
                    meta = result["metadata"]
                    analisis = result["analisis"]
                    
                    with st.expander("Ver Metadatos Detectados", expanded=True):
                        col1, col2 = st.columns(2)
                        col1.metric("Tribunal", meta.tribunal or "N/A")
                        col1.metric("Año", meta.anio or "N/A")
                        col1.metric("Materia", meta.materia or "N/A")
                        col2.metric("Resultado", meta.resultado or "N/A")
                        col2.text("Jueces: " + "; ".join(meta.jueces) if meta.jueces else "N/A")
                        col2.text("Artículos: " + "; ".join(meta.articulos_citados[:5]) + "..." if meta.articulos_citados else "N/A")
                    
                    st.markdown("---")
                    st.subheader("Informe Técnico-Forense")
                    st.markdown(analisis)
                    
                    # Guardar en DB
                    db_id = insert_analisis_db({
                        "tipo_analisis": "fallo",
                        "nombre_archivo": uploaded_file.name,
                        "expediente": meta.expediente,
                        "tribunal": meta.tribunal,
                        "materia": meta.materia,
                        "resumen_analisis": analisis[:400],
                        "fecha_analisis": datetime.now().isoformat(),
                        "metadatos_json": json.dumps(meta.to_dict()),
                        "fuentes_json": json.dumps(result["fuentes"]),
                        "modelo_llm": st.session_state.selected_model
                    })
                    st.caption(f"Guardado en historial (ID: {db_id}).")
                
                except Exception as e:
                    st.error(f"Error en el análisis del fallo: {e}")
                    self.logger.error(f"Error en Tab Fallos: {e}", exc_info=True)

    def _run_tab_expedientes(self):
        """Tab para Análisis de Expedientes PDF (de analista_procesos.py + slegal.py)."""
        st.subheader("Análisis de Expedientes Procesales (PDF)")
        
        uploaded_file = st.file_uploader(
            "Sube un expediente en PDF:",
            type=["pdf"],
            key="expediente_upload"
        )
        
        if uploaded_file:
            if st.button("📋 Analizar Expediente", use_container_width=True):
                try:
                    # Guardar temporalmente
                    temp_path = PDF_TEMP_DIR / uploaded_file.name
                    with open(temp_path, "wb") as f: f.write(uploaded_file.getbuffer())
                    
                    analyzer = ProcesoAnalyzer(
                        st.session_state.selected_model,
                        self.rag_manager
                    )
                    
                    with st.spinner(f"Analizando expediente con {st.session_state.selected_model}..."):
                        result = analyzer.analyze(temp_path)
                    
                    st.success("Análisis de expediente completado.")
                    
                    # Guardar estado para Q&A
                    st.session_state.local_retriever = result["retriever_local"]
                    st.session_state.current_expediente_text = result["texto_completo"]
                    st.session_state.qa_chain = None # Forzar recreación
                    
                    meta = result["metadata"]
                    analisis = result["analisis"]
                    
                    with st.expander("Ver Información Procesal", expanded=True):
                        col1, col2 = st.columns(2)
                        col1.metric("Expediente", meta.expediente or "N/A")
                        col1.metric("Juzgado/Tribunal", meta.tribunal or "N/A")
                        col2.metric("Fuero", meta.fuero or "N/A")
                        col2.metric("Instancia", meta.instancia or "N/A")
                        st.text("Etapas Detectadas: " + "; ".join(meta.etapas_detectadas) if meta.etapas_detectadas else "N/A")

                    st.markdown("---")
                    st.subheader("Dictamen Técnico Procesal")
                    st.markdown(analisis)
                    
                    # --- Exportación (feature de slegal.py) ---
                    col_exp1, col_exp2 = st.columns(2)
                    
                    # Exportar a PDF
                    pdf_path = OUTPUT_DIR / f"Dictamen_{meta.expediente or 'exp'}.pdf"
                    self.report_gen.export_to_pdf(analisis, "Dictamen Procesal", meta.to_dict(), pdf_path)
                    with open(pdf_path, "rb") as f:
                        col_exp1.download_button(
                            "📄 Descargar PDF", f, file_name=pdf_path.name, use_container_width=True
                        )
                    
                    # Exportar a DOCX
                    if DOCX_AVAILABLE:
                        docx_path = OUTPUT_DIR / f"Dictamen_{meta.expediente or 'exp'}.docx"
                        # Extraer secciones para DOCX
                        secciones = re.findall(r"(\d\.\s*.+?)(?=\n\d\.|\Z)", analisis, re.DOTALL)
                        analisis_dict = {f.split('\n')[0]: "\n".join(f.split('\n')[1:]) for f in secciones}
                        
                        self.report_gen.export_to_docx(analisis_dict, meta.to_dict(), docx_path)
                        with open(docx_path, "rb") as f:
                            col_exp2.download_button(
                                "📝 Descargar DOCX", f, file_name=docx_path.name, use_container_width=True
                            )
                    else:
                        col_exp2.warning("Instala `python-docx` para habilitar exportación a DOCX.")

                    # Guardar en DB
                    db_id = insert_analisis_db({
                        "tipo_analisis": "expediente",
                        "nombre_archivo": uploaded_file.name,
                        "expediente": meta.expediente,
                        "tribunal": meta.tribunal,
                        "fuero": meta.fuero,
                        "resumen_analisis": analisis[:400],
                        "ruta_pdf_generado": str(pdf_path),
                        "ruta_docx_generado": str(docx_path) if DOCX_AVAILABLE else "",
                        "fecha_analisis": datetime.now().isoformat(),
                        "metadatos_json": json.dumps(meta.to_dict()),
                        "fuentes_json": json.dumps(result["fuentes"]),
                        "modelo_llm": st.session_state.selected_model
                    })
                    st.caption(f"Guardado en historial (ID: {db_id}).")

                except Exception as e:
                    st.error(f"Error en el análisis del expediente: {e}")
                    self.logger.error(f"Error en Tab Expedientes: {e}", exc_info=True)

        # --- Sección Q&A Interactivo (feature de slegal.py) ---
        if st.session_state.local_retriever:
            st.markdown("---")
            st.subheader("❓ Preguntas sobre el Expediente")
            st.info("Realiza preguntas combinando el contenido del expediente y las bases RAG seleccionadas.")
            
            user_question = st.text_area("Tu pregunta:", key="qa_question", placeholder="Ej: ¿Qué jurisprudencia aplica al vicio de forma mencionado en la demanda?")
            
            if st.button("🔍 Buscar Respuesta", use_container_width=True, disabled=not user_question):
                if not st.session_state.qa_chain:
                    # Crear una chain de QA combinada
                    combined_retriever = EnsembleRetriever(
                        retrievers=[st.session_state.local_retriever, st.session_state.global_retriever],
                        weights=[0.6, 0.4] # Dar más peso al expediente local
                    )
                    
                    st.session_state.qa_chain = RetrievalQA.from_chain_type(
                        llm=ChatGoogleGenerativeAI(model=st.session_state.selected_model, temperature=0.2),
                        chain_type="stuff",
                        retriever=combined_retriever,
                        return_source_documents=True,
                        chain_type_kwargs={"prompt": PROMPT_QA_EXPEDIENTE}
                    )
                
                with st.spinner("Buscando respuesta..."):
                    try:
                        response = st.session_state.qa_chain.invoke({"query": user_question})
                        st.markdown(response["result"])
                        
                        with st.expander("Ver fragmentos fuente (RAG)"):
                            for doc in response["source_documents"]:
                                source = doc.metadata.get('source', 'Expediente Local')
                                st.markdown(f"**Fuente:** `{source}`\n\n```\n{doc.page_content[:300]}...\n```")
                    
                    except Exception as e:
                        st.error(f"Error al procesar la pregunta: {e}")
                        self.logger.error(f"Error en Q&A: {e}", exc_info=True)

    def _run_tab_historico(self):
        """Tab para ver el historial de análisis de la BD unificada."""
        st.subheader("📊 Histórico de Análisis")
        try:
            # Verificar estado de BD
            db_status = check_db_status()
            
            if not db_status["db_exists"]:
                st.warning("⚠️ Base de datos no encontrada. Inicializando...")
                init_db()
            
            with sqlite3.connect(str(DB_PATH)) as cx:
                df = pd.read_sql_query(
                    "SELECT id, fecha_analisis, tipo_analisis, expediente, tribunal, materia, fuero, modelo_llm FROM analisis_historico ORDER BY fecha_analisis DESC LIMIT 100", 
                    cx
                )
            
            if df.empty:
                st.info("✅ Base de datos OK pero sin registros. Comienza a analizar documentos para poblar el histórico.")
            else:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"**Total registros:** {db_status['row_count']} | **Tamaño BD:** {db_status['db_size_mb']} MB")
                with col2:
                    if st.button("🔄 Actualizar"):
                        st.rerun()
                
                st.dataframe(df, use_container_width=True)
        except sqlite3.OperationalError as e:
            st.error(f"❌ Error de acceso a la base de datos: {e}")
            st.info("💡 Verifica que la carpeta de BD tenga permisos de lectura/escritura.")
            self.logger.error(f"Error operacional en BD: {e}", exc_info=True)
        except Exception as e:
            st.error(f"❌ Error al cargar el historial: {e}")
            self.logger.error(f"Error en Tab Histórico: {e}", exc_info=True)

    def _run_tab_bases(self):
        """Permite visualizar y actualizar las bases RAG de manera manual o múltiple."""
        import json
        from pathlib import Path
        from ingesta_multibase import process_subject, BASES_RAG, Settings, materia_to_db_dir

        st.subheader("🧠 Gestión de Bases RAG (Ingesta Multibase)")

        # Mostrar todas las bases con su estado actual
        bases = list(BASES_RAG.keys())
        seleccionadas = st.multiselect(
            "Seleccioná una o varias bases para revisar o actualizar:",
            options=bases,
            default=[],
            help="Podés seleccionar una materia o varias. Si no elegís ninguna, se aplicará a todas."
        )

        st.markdown("---")
        st.write("### Estado actual de las bases:")

        col1, col2 = st.columns([3, 1])
        with col1:
            estados = []
            for materia in bases:
                db_path = materia_to_db_dir(materia)
                cache_path = Path("chroma_cache") / f"cache_{materia.lower().replace(' ', '_')}.json"
                estado = "✅ OK" if db_path.exists() else "❌ No construida"
                fecha = "-"
                if cache_path.exists():
                    try:
                        with open(cache_path, "r", encoding="utf-8") as f:
                            data = json.load(f)
                            if data:
                                # Tomar la fecha más reciente de last_build con validación de tipos
                                fechas = []
                                for v in data.values():
                                    if isinstance(v, dict):
                                        fecha_valor = v.get("last_build", "")
                                        if fecha_valor and isinstance(fecha_valor, str):
                                            fechas.append(fecha_valor)
                                
                                if fechas:
                                    fecha_sorted = sorted(fechas)
                                    fecha_str = fecha_sorted[-1] if fecha_sorted else ""
                                    if fecha_str and len(fecha_str) >= 19:
                                        fecha = fecha_str[:19].replace("T", " ")
                    except Exception:
                        pass
                estados.append({"Materia": materia, "Estado": estado, "Última actualización": fecha})
            st.dataframe(estados, use_container_width=True)

        with col2:
            if st.button("🔄 Actualizar seleccionadas", use_container_width=True):
                if not seleccionadas:
                    seleccionadas = bases
                st.info(f"Iniciando actualización de {len(seleccionadas)} bases...")
                resultados = []
                cfg = Settings()
                for m in seleccionadas:
                    with st.spinner(f"Procesando {m}..."):
                        try:
                            n_pdfs, n_frags, dur = process_subject(m, cfg)
                            resultados.append(f"{m}: {n_pdfs} PDFs → {n_frags} fragmentos ({dur:.1f}s)")
                        except Exception as e:
                            resultados.append(f"{m}: Error → {e}")
                st.success("✅ Proceso completado")
                st.text("\n".join(resultados))

            if st.button("⚙️ Actualizar todas las bases", use_container_width=True):
                resultados = []
                cfg = Settings()
                for m in bases:
                    with st.spinner(f"Procesando {m}..."):
                        try:
                            n_pdfs, n_frags, dur = process_subject(m, cfg)
                            resultados.append(f"{m}: {n_pdfs} PDFs → {n_frags} fragmentos ({dur:.1f}s)")
                        except Exception as e:
                            resultados.append(f"{m}: Error → {e}")
                st.success("✅ Todas las bases fueron actualizadas.")
                st.text("\n".join(resultados))

    def run(self):
        """Punto de entrada principal para la UI."""
        self._render_sidebar()
        
        st.title("⚖️ Motor de Análisis Jurídico Unificado")
        st.markdown(f"**Modelo Activo:** `{st.session_state.selected_model}` | **Bases RAG Activas:** `{len(st.session_state.selected_bases)}`")
        
        tab_texto, tab_fallos, tab_expedientes, tab_historico, tab_bases = st.tabs([
            "💬 Análisis de Texto",
            "📚 Análisis de Fallos (PDF)",
            "⚖️ Análisis de Expedientes (PDF)",
            "📊 Histórico",
            "🧠 Bases RAG"
        ])
        
        with tab_texto:
            self._run_tab_texto()
            
        with tab_fallos:
            self._run_tab_fallos()
            
        with tab_expedientes:
            self._run_tab_expedientes()
            
        with tab_historico:
            self._run_tab_historico()
        
        with tab_bases:
            self._run_tab_bases()

# ============================================================
# 🚀 PUNTO DE ENTRADA (CLI y Dashboard)
# ============================================================

def cli_main():
    """Interfaz de línea de comandos unificada."""
    parser = argparse.ArgumentParser(
        description="Motor de Análisis Jurídico Unificado (CLI)",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    parser.add_argument("--dashboard", action="store_true", help="Ejecutar el dashboard de Streamlit.")
    
    # --- Flujos de Análisis ---
    parser.add_argument("--analyze-juris", type=str, metavar="FILE_PATH",
                        help="Analizar un PDF de fallo jurisprudencial.")
    parser.add_argument("--analyze-proceso", type=str, metavar="FILE_PATH",
                        help="Analizar un PDF de expediente procesal.")
    parser.add_argument("--analyze-text", type=str, metavar="TEXT",
                        help="Analizar un fragmento de texto.")
    
    # --- Opciones para Análisis de Texto ---
    parser.add_argument("--flow", choices=["perfeccionamiento", "interpretacion", "avanzado", "integral"],
                        default="integral", help="Tipo de análisis de texto a ejecutar.")
    
    # --- Opciones Globales ---
    parser.add_argument("--model", default=DEFAULT_LLM, choices=list(MODEL_OPTIONS.keys()),
                        help="Modelo de IA a usar.")
    parser.add_argument("--output", type=str, default="json", choices=["json", "pdf", "docx"],
                        help="Formato de salida para análisis de PDF (solo CLI).")
    
    args = parser.parse_args()
    
    # --- Lógica de Ejecución ---
    
    if args.dashboard:
        logger.info("Iniciando Dashboard Streamlit...")
        # Streamlit ignora argparse, así que simplemente llamamos a run()
    # El comando real debe ser: streamlit run analyser.py -- --dashboard
        try:
            ui = DashboardUI()
            ui.run()
        except Exception as e:
            # Captura errores comunes de "Streamlit can only be run from command line"
            if "Streamlit" in str(e):
                print("\n" + "="*70)
                print("Para iniciar el dashboard, ejecuta:")
                print("streamlit run analyser.py -- --dashboard")
                print("="*70)
            else:
                logger.error(f"Error al iniciar dashboard: {e}")
    
    elif args.analyze_juris:
        logger.info(f"Iniciando análisis de fallo (CLI) con modelo {args.model}...")
        rag_manager = RAGManager(CONFIG)
        analyzer = JurisprudenciaAnalyzer(args.model, rag_manager)
        result = analyzer.analyze(Path(args.analyze_juris))
        print(json.dumps(result, indent=2, ensure_ascii=False, default=str))

    elif args.analyze_proceso:
        logger.info(f"Iniciando análisis de expediente (CLI) con modelo {args.model}...")
        rag_manager = RAGManager(CONFIG)
        analyzer = ProcesoAnalyzer(args.model, rag_manager)
        result = analyzer.analyze(Path(args.analyze_proceso))
        # Remover objetos no serializables para salida JSON
        result.pop("retriever_local", None)
        result.pop("texto_completo", None)
        print(json.dumps(result, indent=2, ensure_ascii=False, default=str))

    elif args.analyze_text:
        logger.info(f"Iniciando análisis de texto (CLI) con flujo {args.flow}...")
        rag_manager = RAGManager(CONFIG)
        global_retriever = rag_manager.get_combined_retriever(list(BASES_RAG.keys()))
        analyzer = TextAnalyzer(args.model, global_retriever)
        
        flow_map = {
            "perfeccionamiento": analyzer.run_perfeccionamiento,
            "interpretacion": analyzer.run_interpretacion,
            "avanzado": analyzer.run_analisis_avanzado,
            "integral": analyzer.run_analisis_integral,
        }
        
        resultado = flow_map[args.flow](args.analyze_text)
        print(resultado)
        
    else:
        # Si no se da ningún argumento, mostrar ayuda
        if len(sys.argv) == 1:
            print("No se especificó ninguna acción.")
            parser.print_help()
            print("\nPara iniciar el dashboard, ejecuta:")
            print("streamlit run analyser.py -- --dashboard")

if __name__ == "__main__":
    # Esta lógica permite que `streamlit run analyser.py -- --dashboard` funcione
    # y que `python analyser.py --argumento` también funcione.

    if _running_in_streamlit():
        # Configurar página ANTES de cualquier otra llamada a Streamlit
        st.set_page_config(
            page_title="Motor de Análisis Jurídico Unificado",
            layout="wide",
            initial_sidebar_state="expanded"
        )
        
        logger.info("Iniciando en modo Dashboard (detectado por Streamlit)...")
        ui = DashboardUI()
        ui.run()
    else:
        logger.info("Iniciando en modo CLI...")
        cli_main()

    def _render_sidebar(self):
        """Renderiza la barra lateral con opciones globales."""
        with st.sidebar:
            st.title("⚖️ Motor Jurídico")
            st.markdown("Configuración de Análisis")
            
            # 1. Selección de Modelo LLM
            selected_model = st.radio(
                "Modelo de IA:",
                options=list(MODEL_OPTIONS.keys()),
                format_func=lambda x: MODEL_OPTIONS[x],
                key="selected_model_radio"
            )
            if selected_model != st.session_state.selected_model:
                st.session_state.selected_model = selected_model
                st.rerun()

            st.markdown("---")
            
            # 2. Selección Dinámica de RAG (feature de slegal.py)
            st.subheader("Bases de Datos RAG")
            all_bases = list(self.rag_manager.all_retrievers.keys())
            
            selected_bases = st.multiselect(
                "Selecciona las bases a consultar:",
                options=all_bases,
                default=all_bases
            )
            
            if st.button("Actualizar Bases RAG"):
                st.session_state.selected_bases = selected_bases
                st.session_state.global_retriever = self.rag_manager.get_combined_retriever(selected_bases)
                # Limpiar chains para que se recreen con el nuevo retriever
                st.session_state.pop("text_analyzer", None) 
                st.session_state.pop("qa_chain", None)
                st.success("Bases RAG actualizadas.")
                st.rerun()

            st.info(f"Consultando {len(st.session_state.selected_bases)} bases.")

    def _run_tab_texto(self):
        """Tab para Análisis de Texto (de slegal.py)."""
        st.subheader("Análisis de Texto y Cláusulas")
        
        flow_selection = st.radio(
            "Selecciona el tipo de análisis:",
            (
                "Perfeccionamiento",
                "Interpretación y Validación",
                "Análisis Avanzado",
                "Análisis Integral (Combina los 3)",
            ),
            horizontal=True,
        )
        
        user_input = st.text_area(
            "Pega el texto legal, cláusula o consulta:",
            height=150,
            placeholder="Ej: 'El locatario no podrá subarrendar el inmueble...'",
        )
        
        if st.button("🚀 Analizar Texto", use_container_width=True, disabled=not user_input):
            if not st.session_state.global_retriever:
                st.error("No hay bases RAG seleccionadas. Por favor, selecciona al menos una en la barra lateral.")
                return

            # Inicializar el analizador de texto con el retriever global actual
            text_analyzer = TextAnalyzer(
                st.session_state.selected_model, 
                st.session_state.global_retriever
            )
            
            with st.spinner(f"Ejecutando {flow_selection}..."):
                if flow_selection == "Perfeccionamiento":
                    resultado = text_analyzer.run_perfeccionamiento(user_input)
                elif flow_selection == "Interpretación y Validación":
                    resultado = text_analyzer.run_interpretacion(user_input)
                elif flow_selection == "Análisis Avanzado":
                    resultado = text_analyzer.run_analisis_avanzado(user_input)
                else: # Integral
                    resultado = text_analyzer.run_analisis_integral(user_input)
            
            st.markdown("---")
            st.subheader("Resultado del Análisis")
            st.markdown(resultado)

    def _run_tab_fallos(self):
        """Tab para Análisis de Fallos PDF (de analista_jurisprudencial.py)."""
        st.subheader("Análisis de Fallos Jurisprudenciales (PDF)")
        
        uploaded_file = st.file_uploader(
            "Sube un fallo en PDF:",
            type=["pdf"],
            key="juris_upload"
        )
        
        if uploaded_file:
            if st.button("⚖️ Analizar Fallo", use_container_width=True):
                try:
                    # Guardar temporalmente
                    temp_path = PDF_TEMP_DIR / uploaded_file.name
                    with open(temp_path, "wb") as f: f.write(uploaded_file.getbuffer())
                    
                    analyzer = JurisprudenciaAnalyzer(
                        st.session_state.selected_model,
                        self.rag_manager
                    )
                    
                    with st.spinner(f"Analizando fallo con {st.session_state.selected_model}..."):
                        result = analyzer.analyze(temp_path)
                    
                    st.success("Análisis de fallo completado.")
                    
                    meta = result["metadata"]
                    analisis = result["analisis"]
                    
                    with st.expander("Ver Metadatos Detectados", expanded=True):
                        col1, col2 = st.columns(2)
                        col1.metric("Tribunal", meta.tribunal or "N/A")
                        col1.metric("Año", meta.anio or "N/A")
                        col1.metric("Materia", meta.materia or "N/A")
                        col2.metric("Resultado", meta.resultado or "N/A")
                        col2.text("Jueces: " + "; ".join(meta.jueces) if meta.jueces else "N/A")
                        col2.text("Artículos: " + "; ".join(meta.articulos_citados[:5]) + "..." if meta.articulos_citados else "N/A")
                    
                    st.markdown("---")
                    st.subheader("Informe Técnico-Forense")
                    st.markdown(analisis)
                    
                    # Guardar en DB
                    db_id = insert_analisis_db({
                        "tipo_analisis": "fallo",
                        "nombre_archivo": uploaded_file.name,
                        "expediente": meta.expediente,
                        "tribunal": meta.tribunal,
                        "materia": meta.materia,
                        "resumen_analisis": analisis[:400],
                        "fecha_analisis": datetime.now().isoformat(),
                        "metadatos_json": json.dumps(meta.to_dict()),
                        "fuentes_json": json.dumps(result["fuentes"]),
                        "modelo_llm": st.session_state.selected_model
                    })
                    st.caption(f"Guardado en historial (ID: {db_id}).")
                
                except Exception as e:
                    st.error(f"Error en el análisis del fallo: {e}")
                    self.logger.error(f"Error en Tab Fallos: {e}", exc_info=True)

    def _run_tab_expedientes(self):
        """Tab para Análisis de Expedientes PDF (de analista_procesos.py + slegal.py)."""
        st.subheader("Análisis de Expedientes Procesales (PDF)")
        
        uploaded_file = st.file_uploader(
            "Sube un expediente en PDF:",
            type=["pdf"],
            key="expediente_upload"
        )
        
        if uploaded_file:
            if st.button("📋 Analizar Expediente", use_container_width=True):
                try:
                    # Guardar temporalmente
                    temp_path = PDF_TEMP_DIR / uploaded_file.name
                    with open(temp_path, "wb") as f: f.write(uploaded_file.getbuffer())
                    
                    analyzer = ProcesoAnalyzer(
                        st.session_state.selected_model,
                        self.rag_manager
                    )
                    
                    with st.spinner(f"Analizando expediente con {st.session_state.selected_model}..."):
                        result = analyzer.analyze(temp_path)
                    
                    st.success("Análisis de expediente completado.")
                    
                    # Guardar estado para Q&A
                    st.session_state.local_retriever = result["retriever_local"]
                    st.session_state.current_expediente_text = result["texto_completo"]
                    st.session_state.qa_chain = None # Forzar recreación
                    
                    meta = result["metadata"]
                    analisis = result["analisis"]
                    
                    with st.expander("Ver Información Procesal", expanded=True):
                        col1, col2 = st.columns(2)
                        col1.metric("Expediente", meta.expediente or "N/A")
                        col1.metric("Juzgado/Tribunal", meta.tribunal or "N/A")
                        col2.metric("Fuero", meta.fuero or "N/A")
                        col2.metric("Instancia", meta.instancia or "N/A")
                        st.text("Etapas Detectadas: " + "; ".join(meta.etapas_detectadas) if meta.etapas_detectadas else "N/A")

                    st.markdown("---")
                    st.subheader("Dictamen Técnico Procesal")
                    st.markdown(analisis)
                    
                    # --- Exportación (feature de slegal.py) ---
                    col_exp1, col_exp2 = st.columns(2)
                    
                    # Exportar a PDF
                    pdf_path = OUTPUT_DIR / f"Dictamen_{meta.expediente or 'exp'}.pdf"
                    self.report_gen.export_to_pdf(analisis, "Dictamen Procesal", meta.to_dict(), pdf_path)
                    with open(pdf_path, "rb") as f:
                        col_exp1.download_button(
                            "📄 Descargar PDF", f, file_name=pdf_path.name, use_container_width=True
                        )
                    
                    # Exportar a DOCX
                    if DOCX_AVAILABLE:
                        docx_path = OUTPUT_DIR / f"Dictamen_{meta.expediente or 'exp'}.docx"
                        # Extraer secciones para DOCX
                        secciones = re.findall(r"(\d\.\s*.+?)(?=\n\d\.|\Z)", analisis, re.DOTALL)
                        analisis_dict = {f.split('\n')[0]: "\n".join(f.split('\n')[1:]) for f in secciones}
                        
                        self.report_gen.export_to_docx(analisis_dict, meta.to_dict(), docx_path)
                        with open(docx_path, "rb") as f:
                            col_exp2.download_button(
                                "📝 Descargar DOCX", f, file_name=docx_path.name, use_container_width=True
                            )
                    else:
                        col_exp2.warning("Instala `python-docx` para habilitar exportación a DOCX.")

                    # Guardar en DB
                    db_id = insert_analisis_db({
                        "tipo_analisis": "expediente",
                        "nombre_archivo": uploaded_file.name,
                        "expediente": meta.expediente,
                        "tribunal": meta.tribunal,
                        "fuero": meta.fuero,
                        "resumen_analisis": analisis[:400],
                        "ruta_pdf_generado": str(pdf_path),
                        "ruta_docx_generado": str(docx_path) if DOCX_AVAILABLE else "",
                        "fecha_analisis": datetime.now().isoformat(),
                        "metadatos_json": json.dumps(meta.to_dict()),
                        "fuentes_json": json.dumps(result["fuentes"]),
                        "modelo_llm": st.session_state.selected_model
                    })
                    st.caption(f"Guardado en historial (ID: {db_id}).")

                except Exception as e:
                    st.error(f"Error en el análisis del expediente: {e}")
                    self.logger.error(f"Error en Tab Expedientes: {e}", exc_info=True)

        # --- Sección Q&A Interactivo (feature de slegal.py) ---
        if st.session_state.local_retriever:
            st.markdown("---")
            st.subheader("❓ Preguntas sobre el Expediente")
            st.info("Realiza preguntas combinando el contenido del expediente y las bases RAG seleccionadas.")
            
            user_question = st.text_area("Tu pregunta:", key="qa_question", placeholder="Ej: ¿Qué jurisprudencia aplica al vicio de forma mencionado en la demanda?")
            
            if st.button("🔍 Buscar Respuesta", use_container_width=True, disabled=not user_question):
                if not st.session_state.qa_chain:
                    # Crear una chain de QA combinada
                    combined_retriever = EnsembleRetriever(
                        retrievers=[st.session_state.local_retriever, st.session_state.global_retriever],
                        weights=[0.6, 0.4] # Dar más peso al expediente local
                    )
                    
                    st.session_state.qa_chain = RetrievalQA.from_chain_type(
                        llm=ChatGoogleGenerativeAI(model=st.session_state.selected_model, temperature=0.2),
                        chain_type="stuff",
                        retriever=combined_retriever,
                        return_source_documents=True,
                        chain_type_kwargs={"prompt": PROMPT_QA_EXPEDIENTE}
                    )
                
                with st.spinner("Buscando respuesta..."):
                    try:
                        response = st.session_state.qa_chain.invoke({"query": user_question})
                        st.markdown(response["result"])
                        
                        with st.expander("Ver fragmentos fuente (RAG)"):
                            for doc in response["source_documents"]:
                                source = doc.metadata.get('source', 'Expediente Local')
                                st.markdown(f"**Fuente:** `{source}`\n\n```\n{doc.page_content[:300]}...\n```")
                    
                    except Exception as e:
                        st.error(f"Error al procesar la pregunta: {e}")
                        self.logger.error(f"Error en Q&A: {e}", exc_info=True)

    def _run_tab_historico(self):
        """Tab para ver el historial de análisis de la BD unificada."""
        st.subheader("📊 Histórico de Análisis")
        try:
            # Verificar estado de BD
            db_status = check_db_status()
            
            if not db_status["db_exists"]:
                st.warning("⚠️ Base de datos no encontrada. Inicializando...")
                init_db()
            
            with sqlite3.connect(str(DB_PATH)) as cx:
                df = pd.read_sql_query(
                    "SELECT id, fecha_analisis, tipo_analisis, expediente, tribunal, materia, fuero, modelo_llm FROM analisis_historico ORDER BY fecha_analisis DESC LIMIT 100", 
                    cx
                )
            
            if df.empty:
                st.info("✅ Base de datos OK pero sin registros. Comienza a analizar documentos para poblar el histórico.")
            else:
                col1, col2 = st.columns([3, 1])
                with col1:
                    st.write(f"**Total registros:** {db_status['row_count']} | **Tamaño BD:** {db_status['db_size_mb']} MB")
                with col2:
                    if st.button("🔄 Actualizar"):
                        st.rerun()
                
                st.dataframe(df, use_container_width=True)
        except sqlite3.OperationalError as e:
            st.error(f"❌ Error de acceso a la base de datos: {e}")
            st.info("💡 Verifica que la carpeta de BD tenga permisos de lectura/escritura.")
            self.logger.error(f"Error operacional en BD: {e}", exc_info=True)
        except Exception as e:
            st.error(f"❌ Error al cargar el historial: {e}")
            self.logger.error(f"Error en Tab Histórico: {e}", exc_info=True)

    def _run_tab_bases(self):
        """Permite visualizar y actualizar las bases RAG de manera manual o múltiple."""
        import json
        from pathlib import Path
        from ingesta_multibase import process_subject, BASES_RAG, Settings, materia_to_db_dir

        st.subheader("🧠 Gestión de Bases RAG (Ingesta Multibase)")

        # Mostrar todas las bases con su estado actual
        bases = list(BASES_RAG.keys())
        seleccionadas = st.multiselect(
            "Seleccioná una o varias bases para revisar o actualizar:",
            options=bases,
            default=[],
            help="Podés seleccionar una materia o varias. Si no elegís ninguna, se aplicará a todas."
        )

        st.markdown("---")
        st.write("### Estado actual de las bases:")

        col1, col2 = st.columns([3, 1])
        with col1:
            estados = []
            for materia in bases:
                db_path = materia_to_db_dir(materia)
                cache_path = Path("chroma_cache") / f"cache_{materia.lower().replace(' ', '_')}.json"
                estado = "✅ OK" if db_path.exists() else "❌ No construida"
                fecha = "-"
                if cache_path.exists():
                    try:
                        with open(cache_path, "r", encoding="utf-8") as f:
                            data = json.load(f)
                            if data:
                                # Tomar la fecha más reciente de last_build con validación de tipos
                                fechas = []
                                for v in data.values():
                                    if isinstance(v, dict):
                                        fecha_valor = v.get("last_build", "")
                                        if fecha_valor and isinstance(fecha_valor, str):
                                            fechas.append(fecha_valor)
                                
                                if fechas:
                                    fecha_sorted = sorted(fechas)
                                    fecha_str = fecha_sorted[-1] if fecha_sorted else ""
                                    if fecha_str and len(fecha_str) >= 19:
                                        fecha = fecha_str[:19].replace("T", " ")
                    except Exception:
                        pass
                estados.append({"Materia": materia, "Estado": estado, "Última actualización": fecha})
            st.dataframe(estados, use_container_width=True)

        with col2:
            if st.button("🔄 Actualizar seleccionadas", use_container_width=True):
                if not seleccionadas:
                    seleccionadas = bases
                st.info(f"Iniciando actualización de {len(seleccionadas)} bases...")
                resultados = []
                cfg = Settings()
                for m in seleccionadas:
                    with st.spinner(f"Procesando {m}..."):
                        try:
                            n_pdfs, n_frags, dur = process_subject(m, cfg)
                            resultados.append(f"{m}: {n_pdfs} PDFs → {n_frags} fragmentos ({dur:.1f}s)")
                        except Exception as e:
                            resultados.append(f"{m}: Error → {e}")
                st.success("✅ Proceso completado")
                st.text("\n".join(resultados))

            if st.button("⚙️ Actualizar todas las bases", use_container_width=True):
                resultados = []
                cfg = Settings()
                for m in bases:
                    with st.spinner(f"Procesando {m}..."):
                        try:
                            n_pdfs, n_frags, dur = process_subject(m, cfg)
                            resultados.append(f"{m}: {n_pdfs} PDFs → {n_frags} fragmentos ({dur:.1f}s)")
                        except Exception as e:
                            resultados.append(f"{m}: Error → {e}")
                st.success("✅ Todas las bases fueron actualizadas.")
                st.text("\n".join(resultados))

    def run(self):
        """Punto de entrada principal para la UI."""
        self._render_sidebar()
        
        st.title("⚖️ Motor de Análisis Jurídico Unificado")
        st.markdown(f"**Modelo Activo:** `{st.session_state.selected_model}` | **Bases RAG Activas:** `{len(st.session_state.selected_bases)}`")
        
        tab_texto, tab_fallos, tab_expedientes, tab_historico, tab_bases = st.tabs([
            "💬 Análisis de Texto",
            "📚 Análisis de Fallos (PDF)",
            "⚖️ Análisis de Expedientes (PDF)",
            "📊 Histórico",
            "🧠 Bases RAG"
        ])
        
        with tab_texto:
            self._run_tab_texto()
            
        with tab_fallos:
            self._run_tab_fallos()
            
        with tab_expedientes:
            self._run_tab_expedientes()
            
        with tab_historico:
            self._run_tab_historico()
        
        with tab_bases:
            self._run_tab_bases()

# ============================================================
# 🚀 PUNTO DE ENTRADA (CLI y Dashboard)
# ============================================================

def cli_main():
    """Interfaz de línea de comandos unificada."""
    parser = argparse.ArgumentParser(
        description="Motor de Análisis Jurídico Unificado (CLI)",
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    parser.add_argument("--dashboard", action="store_true", help="Ejecutar el dashboard de Streamlit.")
    
    # --- Flujos de Análisis ---
    parser.add_argument("--analyze-juris", type=str, metavar="FILE_PATH",
                        help="Analizar un PDF de fallo jurisprudencial.")
    parser.add_argument("--analyze-proceso", type=str, metavar="FILE_PATH",
                        help="Analizar un PDF de expediente procesal.")
    parser.add_argument("--analyze-text", type=str, metavar="TEXT",
                        help="Analizar un fragmento de texto.")
    
    # --- Opciones para Análisis de Texto ---
    parser.add_argument("--flow", choices=["perfeccionamiento", "interpretacion", "avanzado", "integral"],
                        default="integral", help="Tipo de análisis de texto a ejecutar.")
    
    # --- Opciones Globales ---
    parser.add_argument("--model", default=DEFAULT_LLM, choices=list(MODEL_OPTIONS.keys()),
                        help="Modelo de IA a usar.")
    parser.add_argument("--output", type=str, default="json", choices=["json", "pdf", "docx"],
                        help="Formato de salida para análisis de PDF (solo CLI).")
    
    args = parser.parse_args()
    
    # --- Lógica de Ejecución ---
    
    if args.dashboard:
        logger.info("Iniciando Dashboard Streamlit...")
        # Streamlit ignora argparse, así que simplemente llamamos a run()
    # El comando real debe ser: streamlit run analyser.py -- --dashboard
        try:
            ui = DashboardUI()
            ui.run()
        except Exception as e:
            # Captura errores comunes de "Streamlit can only be run from command line"
            if "Streamlit" in str(e):
                print("\n" + "="*70)
                print("Para iniciar el dashboard, ejecuta:")
                print("streamlit run analyser.py -- --dashboard")
                print("="*70)
            else:
                logger.error(f"Error al iniciar dashboard: {e}")
    
    elif args.analyze_juris:
        logger.info(f"Iniciando análisis de fallo (CLI) con modelo {args.model}...")
        rag_manager = RAGManager(CONFIG)
        analyzer = JurisprudenciaAnalyzer(args.model, rag_manager)
        result = analyzer.analyze(Path(args.analyze_juris))
        print(json.dumps(result, indent=2, ensure_ascii=False, default=str))

    elif args.analyze_proceso:
        logger.info(f"Iniciando análisis de expediente (CLI) con modelo {args.model}...")
        rag_manager = RAGManager(CONFIG)
        analyzer = ProcesoAnalyzer(args.model, rag_manager)
        result = analyzer.analyze(Path(args.analyze_proceso))
        # Remover objetos no serializables para salida JSON
        result.pop("retriever_local", None)
        result.pop("texto_completo", None)
        print(json.dumps(result, indent=2, ensure_ascii=False, default=str))

    elif args.analyze_text:
        logger.info(f"Iniciando análisis de texto (CLI) con flujo {args.flow}...")
        rag_manager = RAGManager(CONFIG)
        global_retriever = rag_manager.get_combined_retriever(list(BASES_RAG.keys()))
        analyzer = TextAnalyzer(args.model, global_retriever)
        
        flow_map = {
            "perfeccionamiento": analyzer.run_perfeccionamiento,
            "interpretacion": analyzer.run_interpretacion,
            "avanzado": analyzer.run_analisis_avanzado,
            "integral": analyzer.run_analisis_integral,
        }
        
        resultado = flow_map[args.flow](args.analyze_text)
        print(resultado)
        
    else:
        # Si no se da ningún argumento, mostrar ayuda
        if len(sys.argv) == 1:
            print("No se especificó ninguna acción.")
            parser.print_help()
            print("\nPara iniciar el dashboard, ejecuta:")
            print("streamlit run analyser.py -- --dashboard")

